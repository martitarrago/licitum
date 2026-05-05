"""Renderiza el documento de PROPOSICIÓN ECONÓMICA en HTML y .docx.

Mismo patrón que sobre_a_docx + deuc_generator: HTML para preview en
iframe (con `window.print()` para guardar como PDF) y .docx editable
para que el licitador retoque libremente antes de firmar.

El documento es la pieza formal que va dentro del Sobre Único (en
abierto simplificado) o en el Sobre C (en abierto ordinario). Contiene:
identificación del licitador, expediente, importe ofertado en cifras
y letras, baja %, declaración de aceptación de pliegos, IVA, plazo de
ejecución (si lo modifica el licitador) y firma.
"""
from __future__ import annotations

import io
from html import escape
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


def _e(value: Any) -> str:
    if value is None:
        return ""
    return escape(str(value), quote=True)


def _fmt_eur(v: float | int | None) -> str:
    if v is None:
        return "—"
    try:
        n = float(v)
    except (ValueError, TypeError):
        return "—"
    return f"{n:,.2f} €".replace(",", "X").replace(".", ",").replace("X", ".")


_UNIDADES = ["", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve"]
_DIEZ_A_VEINTE = [
    "diez", "once", "doce", "trece", "catorce", "quince",
    "dieciséis", "diecisiete", "dieciocho", "diecinueve",
]
_DECENAS = [
    "", "", "veinte", "treinta", "cuarenta", "cincuenta",
    "sesenta", "setenta", "ochenta", "noventa",
]
_CENTENAS = [
    "", "ciento", "doscientos", "trescientos", "cuatrocientos",
    "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos",
]


def _num_a_letras_99(n: int) -> str:
    if n < 10:
        return _UNIDADES[n]
    if n < 20:
        return _DIEZ_A_VEINTE[n - 10]
    if n < 30:
        return "veinti" + _UNIDADES[n - 20] if n > 20 else "veinte"
    decena = _DECENAS[n // 10]
    unidad = _UNIDADES[n % 10]
    if unidad:
        return f"{decena} y {unidad}"
    return decena


def _num_a_letras_999(n: int) -> str:
    if n == 100:
        return "cien"
    if n < 100:
        return _num_a_letras_99(n)
    centena = _CENTENAS[n // 100]
    resto = n % 100
    if resto:
        return f"{centena} {_num_a_letras_99(resto)}"
    return centena


def _num_a_letras_int(n: int) -> str:
    """Convierte un entero a letras en castellano. Soporta hasta cientos
    de millones — suficiente para cualquier presupuesto público."""
    if n == 0:
        return "cero"
    partes: list[str] = []
    millones = n // 1_000_000
    miles = (n % 1_000_000) // 1_000
    resto = n % 1_000
    if millones:
        if millones == 1:
            partes.append("un millón")
        else:
            partes.append(f"{_num_a_letras_999(millones)} millones")
    if miles:
        if miles == 1:
            partes.append("mil")
        else:
            partes.append(f"{_num_a_letras_999(miles)} mil")
    if resto:
        partes.append(_num_a_letras_999(resto))
    return " ".join(partes)


def importe_a_letras(importe: float) -> str:
    """Convierte un importe en € a letras. Ej: 487.500,00 → "cuatrocientos
    ochenta y siete mil quinientos euros con cero céntimos"."""
    if importe < 0:
        return "—"
    enteros = int(importe)
    centimos = round((importe - enteros) * 100)
    if centimos >= 100:
        enteros += 1
        centimos -= 100
    enteros_letras = _num_a_letras_int(enteros)
    if centimos == 0:
        return f"{enteros_letras} euros"
    centimos_letras = _num_a_letras_int(centimos)
    return f"{enteros_letras} euros con {centimos_letras} céntimos"


def render_html(snapshot: dict[str, Any]) -> str:
    """HTML de proposición económica para preview en iframe."""
    empresa = snapshot.get("empresa") or {}
    licitacion = snapshot.get("licitacion") or {}
    presupuesto_base = float(snapshot.get("presupuesto_base") or 0)
    baja_pct = float(snapshot.get("baja_pct") or 0)
    importe_ofertado = float(snapshot.get("importe_ofertado") or 0)
    importe_iva = snapshot.get("importe_iva")
    importe_total = snapshot.get("importe_total")
    iva_pct = snapshot.get("iva_pct")
    fecha = snapshot.get("fecha_emision") or ""
    plazo_meses = snapshot.get("plazo_ejecucion_meses")
    nota_riesgo = snapshot.get("nota_riesgo") or ""
    nivel_riesgo = snapshot.get("nivel_riesgo") or "seguro"

    direccion_partes = [
        empresa.get("direccion_calle"),
        empresa.get("direccion_codigo_postal"),
        empresa.get("direccion_ciudad"),
    ]
    direccion = ", ".join(p for p in direccion_partes if p) or ""
    if empresa.get("direccion_provincia"):
        direccion = (
            f"{direccion} ({empresa['direccion_provincia']})"
            if direccion
            else f"({empresa['direccion_provincia']})"
        )

    representante = ""
    if empresa.get("representante_nombre"):
        partes = [empresa["representante_nombre"]]
        if empresa.get("representante_nif"):
            partes.append(f"NIF {empresa['representante_nif']}")
        if empresa.get("representante_cargo"):
            partes.append(empresa["representante_cargo"])
        representante = " · ".join(partes)

    riesgo_banner_html = ""
    if nivel_riesgo == "temerario":
        riesgo_banner_html = (
            '<div class="warning-banner">'
            f"<p><strong>⚠ Aviso interno (no se imprime al firmar):</strong> "
            f"{_e(nota_riesgo)}</p>"
            "</div>"
        )
    elif nivel_riesgo == "atencion":
        riesgo_banner_html = (
            '<div class="info-banner">'
            f"<p><strong>Aviso interno:</strong> {_e(nota_riesgo)}</p>"
            "</div>"
        )

    importe_letras = importe_a_letras(importe_ofertado)

    return f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Proposición económica · {_e(licitacion.get('expediente'))}</title>
<style>
  @page {{ size: A4; margin: 2cm; }}
  body {{
    font-family: "Times New Roman", Times, serif;
    font-size: 11pt;
    line-height: 1.55;
    color: #000;
    background: #fff;
    margin: 0;
    padding: 24pt;
    max-width: 800px;
    margin-inline: auto;
  }}
  h1 {{ font-size: 18pt; margin: 0; letter-spacing: 0.04em; }}
  h2 {{ font-size: 11pt; margin: 1.6em 0 0.4em 0; text-transform: uppercase; letter-spacing: 0.06em; border-bottom: 1px solid #000; padding-bottom: 2pt; }}
  .meta {{ font-size: 9.5pt; color: #444; margin: 0.2em 0 1.6em 0; }}
  .meta strong {{ color: #000; }}
  .section {{ margin-bottom: 1.2em; }}
  dl.datos {{ margin: 0; }}
  dl.datos dt {{ font-weight: bold; width: 32%; float: left; clear: left; }}
  dl.datos dd {{ margin: 0 0 0.25em 32%; }}
  dl.datos::after {{ content: ""; display: block; clear: both; }}
  .importe-box {{
    border: 2px solid #000;
    padding: 1em 1.2em;
    margin: 1em 0;
  }}
  .importe-cifras {{
    font-size: 16pt;
    font-weight: bold;
    margin: 0.3em 0;
  }}
  .importe-letras {{
    font-style: italic;
    color: #333;
    margin: 0.3em 0 0;
  }}
  .baja-pct {{
    font-size: 12pt;
    margin-top: 0.6em;
  }}
  .declaration {{
    margin: 0.5em 0;
    padding-left: 1em;
    border-left: 2px solid #ddd;
    text-align: justify;
  }}
  .signature {{ margin-top: 3em; }}
  .signature-line {{
    margin-top: 5em;
    border-top: 1px solid #000;
    width: 60%;
    padding-top: 4pt;
    font-size: 9.5pt;
  }}
  .warning-banner {{
    background: #fff5e6;
    border-left: 3px solid #d97706;
    padding: 0.6em 1em;
    margin: 1em 0;
    font-size: 9.5pt;
  }}
  .info-banner {{
    background: #f0f7ff;
    border-left: 3px solid #2563eb;
    padding: 0.6em 1em;
    margin: 1em 0;
    font-size: 9.5pt;
  }}
  @media print {{
    body {{ padding: 0; max-width: none; }}
    .warning-banner, .info-banner {{ display: none; }}
  }}
</style>
</head>
<body>

<header>
  <h1>PROPOSICIÓN ECONÓMICA</h1>
  <p class="meta"><em>Sobre Único · oferta económica del licitador</em></p>
  <p class="meta">
    <strong>Expediente:</strong> {_e(licitacion.get('expediente'))}<br>
    {f'<strong>Objeto:</strong> {_e(licitacion.get("titulo"))}<br>' if licitacion.get('titulo') else ''}
    <strong>Órgano contratante:</strong> {_e(licitacion.get('organismo')) or "—"}<br>
    <strong>Fecha:</strong> {_e(fecha)}
  </p>
</header>

{riesgo_banner_html}

<section class="section">
  <h2>Identificación del licitador</h2>
  <dl class="datos">
    <dt>Razón social</dt><dd>{_e(empresa.get('nombre'))}</dd>
    <dt>CIF</dt><dd>{_e(empresa.get('cif'))}</dd>
    {f'<dt>Domicilio</dt><dd>{_e(direccion)}</dd>' if direccion else ''}
    {f'<dt>Representante</dt><dd>{_e(representante)}</dd>' if representante else ''}
  </dl>
</section>

<section class="section">
  <h2>Oferta económica</h2>
  <p>El licitador, conforme al Pliego de Cláusulas Administrativas Particulares y el Pliego de Prescripciones Técnicas que rigen este procedimiento, presenta la siguiente oferta:</p>

  <div class="importe-box">
    <p class="meta" style="margin:0">Importe ofertado (sin IVA)</p>
    <p class="importe-cifras">{_e(_fmt_eur(importe_ofertado))}</p>
    <p class="importe-letras">{_e(importe_letras.capitalize())}.</p>
    <p class="baja-pct">Baja propuesta: <strong>{baja_pct:.2f}% sobre el presupuesto base</strong> ({_e(_fmt_eur(presupuesto_base))}).</p>
  </div>

  {f'<p>IVA aplicable ({iva_pct:.0f}%): {_e(_fmt_eur(importe_iva))}.</p>' if importe_iva is not None else ''}
  {f'<p><strong>Importe total con IVA: {_e(_fmt_eur(importe_total))}.</strong></p>' if importe_total is not None else ''}
  {f'<p>Plazo de ejecución comprometido: <strong>{plazo_meses} meses</strong>.</p>' if plazo_meses else ''}
</section>

<section class="section">
  <h2>Declaraciones</h2>
  <div class="declaration"><strong>1.</strong> El licitador declara <strong>conocer y aceptar íntegramente</strong> el contenido del Pliego de Cláusulas Administrativas Particulares, el Pliego de Prescripciones Técnicas y cuanta documentación rige este procedimiento de contratación.</div>
  <div class="declaration"><strong>2.</strong> La presente oferta económica <strong>incluye todos los gastos</strong> directos e indirectos, beneficio industrial, tributos, tasas y cualquier otro concepto que pueda corresponder, salvo el IVA que se repercute por separado.</div>
  <div class="declaration"><strong>3.</strong> El licitador se <strong>compromete a ejecutar el contrato</strong> conforme a lo declarado, en caso de resultar adjudicatario, dentro del plazo y por el importe aquí ofertados.</div>
</section>

<div class="signature">
  <p>{_e(empresa.get('direccion_ciudad') or '—')}, a {_e(fecha)}</p>
  <div class="signature-line">
    Firma del representante legal &mdash; {_e(empresa.get('representante_nombre') or '...')}
  </div>
</div>

</body>
</html>
"""


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)


def _add_kv(doc: Document, label: str, value: str | None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    val_run = p.add_run(value or "—")
    val_run.font.size = Pt(11)


def _add_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if italic:
        run.italic = True


def render_docx(snapshot: dict[str, Any]) -> bytes:
    """Construye el .docx editable desde la versión guardada."""
    empresa = snapshot.get("empresa") or {}
    licitacion = snapshot.get("licitacion") or {}
    presupuesto_base = float(snapshot.get("presupuesto_base") or 0)
    baja_pct = float(snapshot.get("baja_pct") or 0)
    importe_ofertado = float(snapshot.get("importe_ofertado") or 0)
    importe_iva = snapshot.get("importe_iva")
    importe_total = snapshot.get("importe_total")
    iva_pct = snapshot.get("iva_pct")
    fecha = snapshot.get("fecha_emision") or ""
    plazo_meses = snapshot.get("plazo_ejecucion_meses")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_h1(doc, "PROPOSICIÓN ECONÓMICA")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Sobre Único · oferta económica del licitador")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    _add_h2(doc, "Licitación")
    _add_kv(doc, "Expediente", licitacion.get("expediente"))
    _add_kv(doc, "Objeto", licitacion.get("titulo"))
    _add_kv(doc, "Órgano de contratación", licitacion.get("organismo"))
    if presupuesto_base:
        _add_kv(doc, "Presupuesto base de licitación", _fmt_eur(presupuesto_base))

    _add_h2(doc, "Identificación del licitador")
    _add_kv(doc, "Razón social", empresa.get("nombre"))
    _add_kv(doc, "CIF/NIF", empresa.get("cif"))

    direccion_partes = [
        empresa.get("direccion_calle"),
        empresa.get("direccion_codigo_postal"),
        empresa.get("direccion_ciudad"),
    ]
    direccion = ", ".join(p for p in direccion_partes if p)
    if empresa.get("direccion_provincia"):
        direccion = (
            f"{direccion} ({empresa['direccion_provincia']})"
            if direccion
            else f"({empresa['direccion_provincia']})"
        )
    _add_kv(doc, "Domicilio", direccion or None)
    if empresa.get("representante_nombre"):
        partes = [empresa["representante_nombre"]]
        if empresa.get("representante_nif"):
            partes.append(f"NIF {empresa['representante_nif']}")
        if empresa.get("representante_cargo"):
            partes.append(empresa["representante_cargo"])
        _add_kv(doc, "Representante legal", " · ".join(partes))

    # Bloque importes — muy destacado
    _add_h2(doc, "Oferta económica")
    _add_para(
        doc,
        "El licitador, conforme al Pliego de Cláusulas Administrativas "
        "Particulares y al Pliego de Prescripciones Técnicas, presenta la "
        "siguiente oferta:",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Importe ofertado (sin IVA): ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(_fmt_eur(importe_ofertado))
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(importe_a_letras(importe_ofertado).capitalize() + ".")
    run.italic = True
    run.font.size = Pt(11)

    _add_kv(
        doc,
        "Baja sobre presupuesto base",
        f"{baja_pct:.2f}%  (presupuesto base: {_fmt_eur(presupuesto_base)})",
    )
    if importe_iva is not None and iva_pct is not None:
        _add_kv(doc, f"IVA ({iva_pct:.0f}%)", _fmt_eur(importe_iva))
    if importe_total is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("Importe total con IVA: ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(_fmt_eur(importe_total))
        run.bold = True
        run.font.size = Pt(11)
    if plazo_meses:
        _add_kv(doc, "Plazo de ejecución comprometido", f"{plazo_meses} meses")

    _add_h2(doc, "Declaraciones")
    declaraciones = [
        "El licitador declara conocer y aceptar íntegramente el contenido "
        "del Pliego de Cláusulas Administrativas Particulares, el Pliego "
        "de Prescripciones Técnicas y cuanta documentación rige este "
        "procedimiento de contratación.",
        "La presente oferta económica incluye todos los gastos directos "
        "e indirectos, beneficio industrial, tributos, tasas y cualquier "
        "otro concepto que pueda corresponder, salvo el IVA que se "
        "repercute por separado.",
        "El licitador se compromete a ejecutar el contrato conforme a lo "
        "declarado, en caso de resultar adjudicatario, dentro del plazo "
        "y por el importe aquí ofertados.",
    ]
    for i, decl in enumerate(declaraciones, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(decl)
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ciudad = empresa.get("direccion_ciudad") or "_______________"
    run = p.add_run(f"En {ciudad}, a {fecha}.")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Firma del representante legal")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
