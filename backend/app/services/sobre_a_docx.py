"""Generador del Sobre A en formato .docx (Microsoft Word).

Se construye programáticamente desde el `snapshot` JSONB del registro de
SobreAGeneracion — los mismos datos que alimentan el HTML de preview,
pero con un layout neutral pensado para que el usuario lo edite en Word
antes de imprimir, firmar y subir el PDF al portal.

Diseño deliberadamente sobrio: Calibri 11, sin tipografía editorial.
El usuario va a editar este documento; los floruras se las sumamos al
preview HTML y al render final, no aquí.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

# Declaraciones del art. 159.4 LCSP — el contenido mínimo de la
# declaración responsable en procedimiento abierto simplificado.
# Algunas se rellenan dinámicamente (ej. la 4 sobre solvencia/clasificación)
# y se construyen en `generar_docx`; aquí definimos el resto.
DECLARACIONES_BASE: list[str] = [
    "Ostentar la representación válida y suficiente de la sociedad "
    "licitadora para concurrir al presente procedimiento.",
    "Que la empresa se halla válidamente constituida, dispone de la "
    "capacidad de obrar exigida por el art. 65 LCSP y cuenta con las "
    "autorizaciones administrativas precisas para el ejercicio de la "
    "actividad objeto del contrato.",
    "Que ni la empresa ni sus administradores incurren en ninguna de las "
    "prohibiciones de contratar previstas en el art. 71 LCSP.",
    # La declaración 4 (solvencia/clasificación) se construye dinámicamente.
    "Que la empresa se halla al corriente del cumplimiento de las "
    "obligaciones tributarias con la Hacienda estatal y autonómica "
    "catalana, así como de las obligaciones con la Seguridad Social.",
    "Que la empresa no se encuentra en situación de concurso de acreedores, "
    "declaración de insolvencia, intervención judicial, suspensión de "
    "actividades o disolución.",
    "Que la empresa no ha incurrido en falsedad al emitir declaraciones "
    "responsables o aportar información en procedimientos previos de "
    "contratación pública.",
    "Que el licitador se compromete a adscribir a la ejecución del contrato "
    "los medios personales y materiales suficientes (art. 76.2 LCSP) y a "
    "aportar, en caso de resultar adjudicataria, la documentación "
    "acreditativa de los extremos declarados en el plazo establecido por "
    "la mesa de contratación.",
    # La 9 (designación email) se construye dinámicamente.
    "Que la información y los datos consignados en esta declaración son "
    "ciertos. El firmante conoce que la falsedad podrá ser causa de la "
    "prohibición de contratar prevista en el art. 71.1.e) LCSP.",
]


def _fmt_eur(v: Any) -> str:
    if v is None or v == "":
        return "—"
    try:
        n = float(str(v))
    except (ValueError, TypeError):
        return "—"
    return f"{n:,.0f} €".replace(",", ".")


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)


def _add_kv(doc: Document, label: str, value: str | None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    val_run = p.add_run(value or "—")
    val_run.font.size = Pt(11)


def _add_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    if italic:
        run.italic = True


def generar_docx(snapshot: dict[str, Any]) -> bytes:
    """Construye el .docx desde la versión guardada y devuelve los bytes."""
    empresa = snapshot.get("empresa") or {}
    licitacion = snapshot.get("licitacion") or {}
    usa_relic = bool(snapshot.get("usa_relic"))
    n_registral = snapshot.get("n_registral")
    # `clasificaciones` viene ya filtrada (sólo las relevantes para la
    # licitación: las que matchean la exigida por el PCAP). Si el pliego
    # no exige clasificación, esta lista está vacía y NO se renderiza.
    clasificaciones = snapshot.get("clasificaciones") or []
    clasif_exigida = snapshot.get("clasificacion_exigida")
    docs_extra = snapshot.get("docs_extra") or []
    fecha_emision = snapshot.get("fecha_emision") or ""

    doc = Document()

    # Márgenes A4 amplios
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Estilo base Calibri
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Título ────────────────────────────────────────────────────────
    _add_h1(doc, "DECLARACIÓN RESPONSABLE")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(
        "Documentación administrativa del Sobre A · art. 159.4 LCSP"
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    # ── Datos de la licitación ───────────────────────────────────────
    _add_h2(doc, "Licitación")
    _add_kv(doc, "Expediente", licitacion.get("expediente"))
    _add_kv(doc, "Objeto", licitacion.get("titulo"))
    _add_kv(doc, "Órgano de contratación", licitacion.get("organismo"))
    if licitacion.get("importe_licitacion"):
        _add_kv(
            doc,
            "Presupuesto base de licitación",
            _fmt_eur(licitacion.get("importe_licitacion")),
        )

    # ── Banner RELIC (si aplica) ─────────────────────────────────────
    if usa_relic and n_registral:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("Empresa inscrita en el RELIC. ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(
            f"Conforme al art. 159.4 LCSP, esta inscripción exime al "
            f"licitador de aportar la documentación que ya consta en el "
            f"Registre Electrònic d'Empreses Licitadores i Classificades "
            f"de Catalunya. Número registral: {n_registral}."
        )
        run.font.size = Pt(11)

    # ── Datos del licitador ──────────────────────────────────────────
    _add_h2(doc, "Datos del licitador")
    _add_kv(doc, "Razón social", empresa.get("nombre"))
    _add_kv(doc, "CIF/NIF", empresa.get("cif"))

    direccion_partes = [
        empresa.get("direccion_calle"),
        empresa.get("direccion_codigo_postal"),
        empresa.get("direccion_ciudad"),
    ]
    direccion = ", ".join(p for p in direccion_partes if p)
    if empresa.get("direccion_provincia"):
        direccion = (
            f"{direccion} ({empresa['direccion_provincia']})"
            if direccion
            else f"({empresa['direccion_provincia']})"
        )
    _add_kv(doc, "Domicilio fiscal", direccion or None)

    if empresa.get("telefono"):
        _add_kv(doc, "Teléfono", empresa.get("telefono"))
    if empresa.get("email"):
        _add_kv(doc, "Email habilitado para notificaciones", empresa.get("email"))
    if empresa.get("iae"):
        _add_kv(doc, "Epígrafe IAE", empresa.get("iae"))

    # ── Representante legal ─────────────────────────────────────────
    _add_h2(doc, "Representante legal con poder bastante")
    if empresa.get("representante_nombre"):
        repr_partes = [empresa["representante_nombre"]]
        if empresa.get("representante_nif"):
            repr_partes.append(f"NIF {empresa['representante_nif']}")
        if empresa.get("representante_cargo"):
            repr_partes.append(empresa["representante_cargo"])
        _add_para(doc, " · ".join(repr_partes))
    else:
        _add_para(
            doc,
            "[Pendiente: completar nombre, NIF y cargo del representante en "
            "el módulo Empresa antes de firmar.]",
            italic=True,
        )

    # ── Declaraciones responsables (art. 159.4 LCSP) ─────────────────
    _add_h2(doc, "Declaración responsable (art. 159.4 LCSP)")
    _add_para(
        doc,
        "El representante legal arriba identificado, en nombre y "
        "representación de la empresa licitadora, DECLARA BAJO SU "
        "RESPONSABILIDAD:",
    )

    declaraciones: list[str] = []
    # 1, 2, 3 — base
    declaraciones.extend(DECLARACIONES_BASE[:3])

    # 4 — solvencia/clasificación, dinámico según el pliego
    if clasif_exigida:
        exig_str = f"grupo {clasif_exigida.get('grupo')}"
        if clasif_exigida.get("subgrupo"):
            exig_str += f", subgrupo {clasif_exigida['subgrupo']}"
        if clasif_exigida.get("categoria") is not None:
            exig_str += f", categoría {clasif_exigida['categoria']}"
        if clasificaciones:
            cumplidas = []
            for c in clasificaciones:
                line = c.get("grupo") or ""
                if c.get("subgrupo"):
                    line += f"-{c['subgrupo']}"
                if c.get("categoria") is not None:
                    line += f", cat. {c['categoria']}"
                if c.get("fuente") == "relic":
                    line += " (RELIC)"
                cumplidas.append(line)
            decl4 = (
                "Que la empresa cumple los requisitos de solvencia económica "
                "y financiera y técnica o profesional exigidos por el PCAP. "
                f"Clasificación exigida: {exig_str}. "
                f"Clasificación acreditada: {'; '.join(cumplidas)}."
            )
        else:
            decl4 = (
                "Que la empresa cumple los requisitos de solvencia económica "
                "y financiera y técnica o profesional exigidos por el PCAP. "
                f"[ATENCIÓN: el PCAP exige clasificación {exig_str} pero esta "
                "no consta vigente en los datos del licitador — revisar "
                "antes de firmar.]"
            )
    else:
        decl4 = (
            "Que la empresa cumple los requisitos de solvencia económica y "
            "financiera y técnica o profesional exigidos por el Pliego de "
            "Cláusulas Administrativas Particulares."
        )
    declaraciones.append(decl4)

    # 5, 6, 7, 8 — base (índices 3, 4, 5, 6 en DECLARACIONES_BASE)
    declaraciones.extend(DECLARACIONES_BASE[3:7])

    # 9 — designación email para notificaciones, dinámico
    email = empresa.get("email")
    if email:
        decl9 = (
            f"Que el licitador designa la dirección de correo electrónico "
            f"{email} como medio preferente para la práctica de las "
            f"notificaciones derivadas de este procedimiento."
        )
    else:
        decl9 = (
            "Que el licitador designa como medio preferente para la práctica "
            "de las notificaciones la dirección de correo electrónico "
            "[completar antes de firmar]."
        )
    declaraciones.append(decl9)

    # 10 — veracidad (último elemento de DECLARACIONES_BASE)
    declaraciones.append(DECLARACIONES_BASE[7])

    for i, decl in enumerate(declaraciones, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(decl)
        run.font.size = Pt(11)

    # ── Documentación adicional exigida por el pliego ────────────────
    if docs_extra:
        _add_h2(doc, "Documentación adicional exigida por el pliego")
        _add_para(
            doc,
            "Se aporta la documentación complementaria requerida en el PCAP:",
        )
        for d in docs_extra:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run("• ")
            run.font.size = Pt(11)
            run = p.add_run(d)
            run.font.size = Pt(11)

    # ── Lugar y fecha + firma ───────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ciudad = empresa.get("direccion_ciudad") or "_______________"
    run = p.add_run(f"En {ciudad}, a {fecha_emision}.")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Firma del representante legal")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    # ── Serializar a bytes ──────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
